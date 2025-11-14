"""
Document Generator Service
Generates Word documents and PDFs from resume data
"""
import subprocess
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Optional
import tempfile
from app.config.logger import get_logger

logger = get_logger("document_generator")

class DocumentGenerator:
    """Generate Word and PDF documents from resume data"""
    
    # Color scheme for professional resume
    PRIMARY_COLOR = RGBColor(25, 118, 210)  # Blue
    SECONDARY_COLOR = RGBColor(66, 66, 66)  # Dark gray
    TEXT_COLOR = RGBColor(0, 0, 0)  # Black
    
    @staticmethod
    def generate_docx(resume_data: Dict, output_path: str) -> str:
        """
        Generate Word document from resume data
        
        Args:
            resume_data: Complete resume data dictionary
            output_path: Path where document will be saved
            
        Returns:
            Path to generated document
        """
        
        try:
            doc = Document()
            
            # Set up margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            personal_info = resume_data.get('personal_info', {})
            
            # Header - Name
            name_para = doc.add_paragraph()
            name_run = name_para.add_run(personal_info.get('name', 'Your Name'))
            name_run.font.size = Pt(24)
            name_run.font.bold = True
            name_run.font.color.rgb = DocumentGenerator.PRIMARY_COLOR
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Contact Information
            contact_para = doc.add_paragraph()
            contact_text = f"{personal_info.get('email', '')} | {personal_info.get('phone', '')}"
            if personal_info.get('location'):
                contact_text += f" | {personal_info.get('location')}"
            if personal_info.get('website'):
                contact_text += f" | {personal_info.get('website')}"
            
            contact_run = contact_para.add_run(contact_text)
            contact_run.font.size = Pt(10)
            contact_run.font.color.rgb = DocumentGenerator.SECONDARY_COLOR
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Spacing
            
            # Professional Summary
            if resume_data.get('summary'):
                DocumentGenerator._add_section(doc, 'Professional Summary')
                summary_para = doc.add_paragraph(resume_data['summary'])
                summary_para.paragraph_format.space_after = Pt(12)
            
            # Professional Experience
            if resume_data.get('experience') and len(resume_data['experience']) > 0:
                DocumentGenerator._add_section(doc, 'Professional Experience')
                
                for exp in resume_data['experience']:
                    # Position and Company
                    exp_para = doc.add_paragraph()
                    exp_run = exp_para.add_run(
                        f"{exp.get('position', '')} at {exp.get('company', '')}"
                    )
                    exp_run.font.bold = True
                    exp_run.font.size = Pt(12)
                    exp_run.font.color.rgb = DocumentGenerator.PRIMARY_COLOR
                    
                    # Dates and Location
                    date_para = doc.add_paragraph()
                    date_text = f"{exp.get('start_date', '')} - {exp.get('end_date', '')}"
                    if exp.get('location'):
                        date_text += f" | {exp.get('location')}"
                    date_para.add_run(date_text)
                    date_para.paragraph_format.space_before = Pt(0)
                    date_para.paragraph_format.space_after = Pt(6)
                    
                    # Responsibilities as bullet points
                    for responsibility in exp.get('responsibilities', []):
                        doc.add_paragraph(responsibility, style='List Bullet')
                    
                    doc.add_paragraph()  # Spacing between experiences
            
            # Education
            if resume_data.get('education') and len(resume_data['education']) > 0:
                DocumentGenerator._add_section(doc, 'Education')
                
                for edu in resume_data['education']:
                    # Degree and Field
                    edu_para = doc.add_paragraph()
                    edu_run = edu_para.add_run(
                        f"{edu.get('degree', '')} in {edu.get('field', '')}"
                    )
                    edu_run.font.bold = True
                    edu_run.font.color.rgb = DocumentGenerator.PRIMARY_COLOR
                    
                    # Institution and Dates
                    school_para = doc.add_paragraph()
                    school_text = edu.get('institution', '')
                    if edu.get('start_date') or edu.get('end_date'):
                        school_text += f" | {edu.get('start_date', '')} - {edu.get('end_date', '')}"
                    if edu.get('gpa'):
                        school_text += f" | GPA: {edu.get('gpa')}"
                    school_para.add_run(school_text)
                    school_para.paragraph_format.space_before = Pt(0)
                    school_para.paragraph_format.space_after = Pt(12)
            
            # Skills
            if resume_data.get('skills') and len(resume_data['skills']) > 0:
                DocumentGenerator._add_section(doc, 'Skills')
                skills_text = ', '.join(resume_data['skills'])
                skills_para = doc.add_paragraph(skills_text)
                skills_para.paragraph_format.space_after = Pt(12)
            
            # Projects
            if resume_data.get('projects') and len(resume_data['projects']) > 0:
                DocumentGenerator._add_section(doc, 'Projects')
                
                for project in resume_data['projects']:
                    # Project Name
                    proj_para = doc.add_paragraph()
                    proj_run = proj_para.add_run(project.get('name', ''))
                    proj_run.font.bold = True
                    proj_run.font.color.rgb = DocumentGenerator.PRIMARY_COLOR
                    
                    # Description
                    doc.add_paragraph(project.get('description', ''))
                    
                    # Technologies
                    if project.get('technologies'):
                        tech_text = f"Technologies: {', '.join(project.get('technologies', []))}"
                        tech_para = doc.add_paragraph(tech_text)
                        tech_para.paragraph_format.space_after = Pt(12)
                    
                    # Project Link
                    if project.get('link'):
                        link_para = doc.add_paragraph()
                        link_run = link_para.add_run(project.get('link'))
                        link_run.font.italic = True
                        link_para.paragraph_format.space_after = Pt(12)
            
            # Certifications
            if resume_data.get('certifications') and len(resume_data['certifications']) > 0:
                DocumentGenerator._add_section(doc, 'Certifications')
                
                for cert in resume_data['certifications']:
                    doc.add_paragraph(cert, style='List Bullet')
            
            # Save document
            doc.save(output_path)
            logger.info(f"Word document generated: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating Word document: {str(e)}")
            raise
    
    @staticmethod
    def _add_section(doc: Document, section_title: str) -> None:
        """
        Add a formatted section header to document
        
        Args:
            doc: Document object
            section_title: Title of the section
        """
        # Section heading
        heading = doc.add_paragraph()
        heading_run = heading.add_run(section_title.upper())
        heading_run.font.size = Pt(13)
        heading_run.font.bold = True
        heading_run.font.color.rgb = DocumentGenerator.PRIMARY_COLOR
        
        # Add line after heading
        heading.paragraph_format.space_after = Pt(6)
        
        # Add horizontal line effect using table (optional)
        pPr = heading._element.get_or_add_pPr()
        pBdr = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr')
        if pBdr is None:
            from docx.oxml import OxmlElement
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'single')
            bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz', '12')
            bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}space', '1')
            bottom.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color', '1976D2')
            pBdr.append(bottom)
            pPr.append(pBdr)
    
    @staticmethod
    def generate_pdf_from_docx(docx_path: str, pdf_path: str) -> Optional[str]:
        """
        Convert DOCX to PDF using LibreOffice
        
        Note: Requires LibreOffice to be installed on the system
        
        Args:
            docx_path: Path to input DOCX file
            pdf_path: Path where PDF will be saved
            
        Returns:
            Path to PDF if successful, None otherwise
        """
        try:
            # Using LibreOffice for conversion
            subprocess.run([
                'libreoffice',
                '--headless',
                '--convert-to',
                'pdf',
                '--outdir',
                os.path.dirname(pdf_path),
                docx_path
            ], check=True, timeout=60, capture_output=True)
            
            logger.info(f"PDF generated: {pdf_path}")
            return pdf_path
            
        except FileNotFoundError:
            logger.error("LibreOffice not found. Please install it for PDF conversion.")
            logger.warning("PDF conversion skipped. User can convert DOCX to PDF manually.")
            return None
        except subprocess.TimeoutExpired:
            logger.error("PDF conversion timed out")
            return None
        except Exception as e:
            logger.error(f"Error converting to PDF: {str(e)}")
            return None
    
    @staticmethod
    def generate_from_latex(latex_content: str, output_path: str) -> Optional[str]:
        """
        Generate PDF from LaTeX content
        
        Note: Requires pdflatex to be installed
        
        Args:
            latex_content: LaTeX document content
            output_path: Path where PDF will be saved
            
        Returns:
            Path to PDF if successful, None otherwise
        """
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                tex_file = os.path.join(tmpdir, 'resume.tex')
                
                # Write LaTeX content to file
                with open(tex_file, 'w', encoding='utf-8') as f:
                    f.write(latex_content)
                
                logger.info("Compiling LaTeX to PDF...")
                
                # Compile LaTeX
                subprocess.run([
                    'pdflatex',
                    '-output-directory', tmpdir,
                    '-interaction=nonstopmode',
                    tex_file
                ], check=True, timeout=60, capture_output=True)
                
                pdf_file = os.path.join(tmpdir, 'resume.pdf')
                
                if os.path.exists(pdf_file):
                    # Copy to output path
                    import shutil
                    shutil.copy(pdf_file, output_path)
                    logger.info(f"PDF from LaTeX generated: {output_path}")
                    return output_path
                else:
                    logger.error("PDF file not created after LaTeX compilation")
                    return None
                    
        except FileNotFoundError:
            logger.error("pdflatex not found. Please install LaTeX distribution.")
            return None
        except subprocess.TimeoutExpired:
            logger.error("LaTeX compilation timed out")
            return None
        except Exception as e:
            logger.error(f"Error generating PDF from LaTeX: {str(e)}")
            return None
